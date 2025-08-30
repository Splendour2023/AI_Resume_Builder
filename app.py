from flask import Flask, render_template, request, send_file, session
from groq import Groq
import os
from dotenv import load_dotenv
import markdown
from io import BytesIO
from docx import Document
import pdfkit

# -------------------------------
# Load environment variables
# -------------------------------
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
SECRET_KEY = os.getenv("SECRET_KEY", "supersecret")
DEFAULT_MODEL = os.getenv("DEFAULT_MODEL", "openai/gpt-oss-20b")

# -------------------------------
# Initialize Flask and Groq client
# -------------------------------
app = Flask(__name__)
app.secret_key = SECRET_KEY
client = Groq(api_key=GROQ_API_KEY)

# -------------------------------
# Routes
# -------------------------------
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Required fields
        name = request.form.get("name", "").strip()
        role = request.form.get("role", "").strip()
        skills = request.form.get("skills", "").strip()

        if not name or not role or not skills:
            return "Please provide Name, Role, and Skills.", 400

        # Optional fields
        company = request.form.get("company", "the company").strip()
        experience = request.form.get("experience_text", "").strip()
        email = request.form.get("email", "").strip()
        phone = request.form.get("phone", "").strip()
        linkedin = request.form.get("linkedin", "").strip()
        github = request.form.get("github", "").strip()
        address = request.form.get("address", "").strip()

        # Contact info formatting
        contact_info = ""
        if email: contact_info += f"📧 Email: {email}\n"
        if phone: contact_info += f"📞 Phone: {phone}\n"
        if linkedin: contact_info += f"💼 LinkedIn: {linkedin}\n"
        if github: contact_info += f"💻 GitHub/Portfolio: {github}\n"
        if address: contact_info += f"📍 Address: {address}\n"

        # Prompts
        resume_prompt = f"""
Create a professional Markdown resume for {name} applying for a {role} role at {company}.
Skills: {skills}.
Experience summary: {experience if experience else 'N/A'}.
Contact:
{contact_info if contact_info else 'No contact info provided.'}
"""

        cover_prompt = f"""
Write a professional Markdown cover letter for {name} applying for a {role} role at {company}.
Skills: {skills}.
Experience summary: {experience if experience else 'N/A'}.
Contact:
{contact_info if contact_info else 'No contact info provided.'}
"""

        # Call Groq API safely
        try:
            resume_response = client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[{"role": "user", "content": resume_prompt}],
                max_tokens=600
            )
            cover_response = client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[{"role": "user", "content": cover_prompt}],
                max_tokens=500
            )
        except Exception as e:
            return f"<h2>Error from AI model:</h2><pre>{e}</pre>"

        # Extract text safely
        resume_md = getattr(resume_response.choices[0].message, "content", "").strip()
        cover_md = getattr(cover_response.choices[0].message, "content", "").strip()

        if not resume_md: resume_md = "No resume content generated."
        if not cover_md: cover_md = "No cover letter content generated."

        # Save to session
        session["resume"] = resume_md
        session["cover_letter"] = cover_md

        # Convert to HTML
        resume_html = markdown.markdown(resume_md)
        cover_html = markdown.markdown(cover_md)

        return render_template("result.html", resume=resume_html, cover_letter=cover_html)

    return render_template("index.html")

# -------------------------------
# Download PDF
# -------------------------------
@app.route("/download/pdf")
def download_pdf():
    resume_html = markdown.markdown(session.get("resume", ""))
    cover_html = markdown.markdown(session.get("cover_letter", ""))

    full_html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
      <meta charset="utf-8">
      <style>
        body {{ font-family: Arial, sans-serif; background:#f4f6f9; color:#222; margin:20px; line-height:1.5; }}
        h1 {{ color:#0077b6; margin-bottom:10px; }}
        h2 {{ color:#00b4d8; border-bottom:1px solid #0077b6; padding-bottom:4px; }}
        .section {{ margin-bottom:25px; }}
        .resume, .cover {{ background:#fbfdff; padding:16px; border-radius:8px; border:1px solid #eef7ff; }}
        p {{ margin:4px 0; }}
      </style>
    </head>
    <body>
      <div class="section"><h1>Resume</h1><div class="resume">{resume_html}</div></div>
      <div class="section"><h1>Cover Letter</h1><div class="cover">{cover_html}</div></div>
    </body>
    </html>
    """

    pdf_bytes = pdfkit.from_string(full_html, False, options={"encoding": "UTF-8"})
    return send_file(BytesIO(pdf_bytes), as_attachment=True, download_name="resume_coverletter.pdf", mimetype="application/pdf")

# -------------------------------
# Download Word
# -------------------------------
@app.route("/download/word")
def download_word():
    resume = session.get("resume", "")
    cover_letter = session.get("cover_letter", "")
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Resume", level=1)
    doc.add_paragraph(resume)
    doc.add_heading("Cover Letter", level=1)
    doc.add_paragraph(cover_letter)
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="resume_coverletter.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# -------------------------------
# Download TXT
# -------------------------------
@app.route("/download/txt")
def download_txt():
    resume = session.get("resume", "")
    cover_letter = session.get("cover_letter", "")
    text_content = f"=== Resume ===\n{resume}\n\n=== Cover Letter ===\n{cover_letter}"
    buffer = BytesIO(text_content.encode("utf-8"))
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="resume_coverletter.txt", mimetype="text/plain")

# -------------------------------
# Download HTML
# -------------------------------
@app.route("/download/html")
def download_html():
    resume_html = markdown.markdown(session.get("resume", ""))
    cover_html = markdown.markdown(session.get("cover_letter", ""))
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head><meta charset="utf-8"></head>
    <body>
        <h1>Resume</h1>{resume_html}
        <h1>Cover Letter</h1>{cover_html}
    </body>
    </html>
    """
    buffer = BytesIO(html_content.encode("utf-8"))
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="resume_coverletter.html", mimetype="text/html")

# -------------------------------
# Download Image (Header)
# -------------------------------
@app.route("/download/image")
def download_image():
    image_path = "static/resume_header.jpg"  # change to PNG if needed
    return send_file(image_path, as_attachment=True, download_name="resume_header.jpg", mimetype="image/jpeg")

# -------------------------------
# Run Flask
# -------------------------------
if __name__ == "__main__":
    app.run(debug=True)